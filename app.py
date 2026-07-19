import os
import tempfile
import warnings
from datetime import datetime
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS  # مهمة جداً للسماح بالاتصال من موقعك
import yt_dlp
import whisper
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

warnings.filterwarnings("ignore")

app = Flask(__name__)
CORS(app)  # تفعيل الاتصال المفتوح الآمن

OUTPUT_DIR = tempfile.gettempdir()

def download_video(url, output_path):
    try:
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': output_path,
            'quiet': True,
            'no_warnings': True,
            # خيارات إضافية لتجاوز حظر يوتيوب على السيرفرات
            'nocheckcertificate': True,
            'ignoreerrors': True,
            'no_color': True,
            'geo_bypass': True,
            'source_address': '0.0.0.0',
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9',
                'Sec-Fetch-Mode': 'navigate'
            }
        }
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])
        return True
    except Exception as e:
        return False

def transcribe_audio(file_path, model_size="medium"):
    try:
        model = whisper.load_model(model_size)
        result = model.transcribe(
            file_path, 
            language="ar", 
            fp16=False,
            condition_on_previous_text=False,
            no_speech_threshold=0.6 
        )
        return result["text"]
    except Exception as e:
        return None

def save_to_word(text, filepath):
    doc = Document()
    heading = doc.add_heading('نص الدرس المحوّل / Transcribed Text', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p_date = doc.add_paragraph(f'Date/التاريخ: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p_line = doc.add_paragraph('_' * 50)
    p_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    paragraphs = text.split('.')
    for p in paragraphs:
        if p.strip():
            paragraph = doc.add_paragraph(p.strip() + '.')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.right_to_left = True
    doc.save(filepath)

@app.route('/process', methods=['POST'])
def process_video():
    data = request.json
    url = data.get('url')
    model_size = data.get('model', 'medium')
    
    if not url:
        return jsonify({'success': False, 'error_ar': 'الرجاء إدخال الرابط!', 'error_en': 'Please enter a URL!'}), 400

    temp_dir = tempfile.gettempdir()
    audio_path = os.path.join(temp_dir, f"web_{datetime.now().strftime('%H%M%S')}.m4a")
    
    if not download_video(url, audio_path):
        return jsonify({'success': False, 'error_ar': 'فشل تحميل الفيديو.', 'error_en': 'Failed to download video.'}), 500

    text = transcribe_audio(audio_path, model_size)
    
    if os.path.exists(audio_path):
        os.remove(audio_path)

    if not text:
        return jsonify({'success': False, 'error_ar': 'فشل تحويل الصوت.', 'error_en': 'Transcription failed.'}), 500

    filename = f"Transcribed_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc_path = os.path.join(OUTPUT_DIR, filename)
    save_to_word(text, doc_path)

    return jsonify({
        'success': True,
        'preview': text[:300] + "...",
        'file_id': filename
    })

@app.route('/download/<filename>')
def download_file(filename):
    file_path = os.path.join(OUTPUT_DIR, filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return "File not found", 404

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
